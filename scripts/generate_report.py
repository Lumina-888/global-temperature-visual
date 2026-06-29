"""
generate_report.py — 项目答辩 Word 文档生成器
================================================
这个脚本用 python-docx 库自动生成一份专业的 Word 文档，包含：
  - 封面（标题、副标题、技术栈说明）
  - 目录
  - 项目概述
  - 数据采集与处理流程
  - 系统架构设计
  - 后端模块说明
  - 前端可视化模块
  - 关键技术难点与解决方案
  - 项目总结
  - 附录（文件清单）

运行方式：
    cd scripts/
    python generate_report.py

输出：
    项目根目录/答辩报告_全球气温可视化.docx

python-docx 基础概念：
  Document  = 整个 Word 文档
  Paragraph = 一段文字
  Run       = 一段文字中的"连续同格式片段"
              比如一句话里有两个字加粗了，就是 3 个 Run
  Table     = 表格
"""

from pathlib import Path

from docx import Document                          # Word 文档对象
from docx.shared import Cm, Inches, Pt, RGBColor   # 尺寸和颜色单位
from docx.enum.text import WD_ALIGN_PARAGRAPH      # 段落对齐方式
from docx.enum.table import WD_TABLE_ALIGNMENT     # 表格对齐方式
from docx.oxml.ns import qn                        # Word 内部 XML 命名空间工具


# ═══════════════════════════════════════════════════════════════
# 工具函数 — 避免重复代码
# ═══════════════════════════════════════════════════════════════

def add_heading_styled(doc, text, level):
    """
    添加有样式的标题（深蓝色字体）。

    python-docx 的 add_heading() 会创建一个带默认样式的段落。
    但默认颜色是黑的，我们需要改成深蓝(#1A3C6E)来匹配项目主题色。

    Args:
        doc:   Document 对象
        text:  标题文字
        level: 标题级别（1=一级标题, 2=二级标题, ...）

    Returns:
        创建的标题段落对象
    """
    h = doc.add_heading(text, level=level)
    # add_heading 返回的段落里可能包含多个 run（格式片段）
    # 遍历所有 run，统一设成深蓝色
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
    return h


def add_code_block(doc, code_text):
    """
    添加代码块（灰色背景 + Consolas 字体）。

    在 Word 里做代码块比较麻烦——没有原生的"代码块"功能。
    需要用底纹（shading）模拟灰色背景。

    实现原理：
      给段落的 run 添加 XML 属性 w:shd（Word 的底纹元素），
      设置 fill 颜色为 #F0F0F0（浅灰）。

    Args:
        doc:       Document 对象
        code_text: 代码文字（可以包含 \n 换行）

    Returns:
        创建的段落对象
    """
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)      # 左边缩进半厘米
    p.paragraph_format.space_before = Pt(2)        # 段前间距 2pt（紧凑）
    p.paragraph_format.space_after = Pt(2)         # 段后间距 2pt

    run = p.add_run(code_text)
    run.font.name = "Consolas"                      # 等宽字体，代码看起来整齐
    run.font.size = Pt(8)                           # 8pt 小字——代码块不需要太大
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33) # 深灰文字

    # ── 添加灰色底纹 ──
    # python-docx 没有直接给 run 设置背景色的 API，
    # 需要手动操作底层 XML 元素（lxml）
    shd = run._element.rPr  # rPr = Run Properties（run 的属性容器）
    if shd is None:
        from lxml import etree
        shd = etree.SubElement(run._element, qn('w:rPr'))

    shd_elem = shd.find(qn('w:shd'))  # shd = Shading（底纹）
    if shd_elem is None:
        from lxml import etree
        shd_elem = etree.SubElement(shd, qn('w:shd'))

    shd_elem.set(qn('w:fill'), 'F0F0F0')  # 填充色 = 浅灰
    shd_elem.set(qn('w:val'), 'clear')     # 底纹类型 = 纯色填充

    return p


def set_cell_shading(cell, color):
    """
    给表格单元格设置背景色。

    同样需要操作底层 XML——python-docx 的表格 API 比较简单。

    Args:
        cell:  单元格对象
        color: 颜色值（如 '1A3C6E' 深蓝、'EBF0F7' 浅蓝灰）
    """
    shd = cell._element.get_or_add_tcPr()  # tcPr = Table Cell Properties
    from lxml import etree
    shd_elem = shd.find(qn('w:shd'))
    if shd_elem is None:
        shd_elem = etree.SubElement(shd, qn('w:shd'))
    shd_elem.set(qn('w:fill'), color)
    shd_elem.set(qn('w:val'), 'clear')


def add_table_with_style(doc, headers, rows, col_widths=None):
    """
    添加带样式的表格。

    表格样式：
      - 表头：深蓝背景、白色粗体文字、居中
      - 数据行：交替背景色（斑马纹）、8.5pt 小字
      - 可选列宽设置

    Args:
        doc:        Document 对象
        headers:    表头列表，如 ["姓名", "年龄", "城市"]
        rows:       数据行列表，如 [["张三", "25", "北京"], ["李四", "30", "上海"]]
        col_widths: 可选，每列的宽度（厘米），如 [3.0, 2.0, 5.0]

    Returns:
        创建的表格对象
    """
    # doc.add_table(行数, 列数)
    # 行数 = 1(表头) + 数据行数
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'  # 使用 Word 内置的表格样式（带边框）
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # ── 表头 ──
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 水平居中
            for run in p.runs:
                run.font.bold = True                  # 加粗
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)  # 白色字
        set_cell_shading(cell, '1A3C6E')  # 深蓝背景

    # ── 数据行 ──
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]  # r+1 跳过表头行
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(8.5)
            # 斑马纹：偶数行（r % 2 == 0）加浅蓝灰背景
            if r % 2 == 0:
                set_cell_shading(cell, 'EBF0F7')

    # ── 设置列宽（可选）──
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph()  # 表格后面加个空行，防止和下段文字贴太紧
    return table


# ═══════════════════════════════════════════════════════════════
# 主函数 — 生成完整文档
# ═══════════════════════════════════════════════════════════════

def main():
    # ═══════════════════════════════════════════════════════════
    # 创建文档 + 设置默认样式
    # ═══════════════════════════════════════════════════════════
    doc = Document()

    # 设置 Normal 样式（= 正文默认样式）
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'          # 西文字体
    style.font.size = Pt(10.5)           # 五号字
    style.paragraph_format.line_spacing = 1.5  # 1.5 倍行距
    style.paragraph_format.space_after = Pt(4) # 段后 4pt 间距
    # 设置中文字体（Word 里中西文字体是分开的）
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # 页面设置（A4 纸标准）
    section = doc.sections[0]
    section.page_width = Cm(21)          # A4 宽
    section.page_height = Cm(29.7)       # A4 高
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    # ═══════════════════════════════════════════════════════════
    # 封面
    # ═══════════════════════════════════════════════════════════
    # 用 6 个空行把标题推到页面中间偏上

    for _ in range(6):
        doc.add_paragraph()  # 空行

    # 主标题
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 居中
    run = title.add_run("全球气温变化可视化分析系统")
    run.font.size = Pt(26)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)  # 深蓝

    # 副标题
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("—— 基于 NASA GISTEMP v4 与 Berkeley Earth 数据集")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)  # 中灰

    doc.add_paragraph()
    doc.add_paragraph()

    # 项目信息
    info_lines = [
        "技术栈：Python FastAPI + ECharts + Pandas",
        "数据来源：NASA GISTEMP v4 / GHCN v4 台站数据 / Berkeley Earth",
        "数据范围：1750 年 — 2026 年（全球） | 1750 年 — 2024 年（城市/国家）",
    ]
    for line in info_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_page_break()  # 封面结束，换新页

    # ═══════════════════════════════════════════════════════════
    # 目录页
    # ═══════════════════════════════════════════════════════════
    add_heading_styled(doc, "目  录", 1)

    toc_items = [
        "一、项目概述",
        "二、数据采集与处理",
        "  2.1 数据来源介绍",
        "  2.2 NASA GISTEMP 数据采集流程",
        "  2.3 数据格式转换与导入",
        "三、系统架构设计",
        "  3.1 整体架构",
        "  3.2 技术栈",
        "四、后端模块说明",
        "  4.1 FastAPI 路由层 (app.py)",
        "  4.2 数据引擎层 (data_engine.py)",
        "  4.3 数据更新脚本 (update_data.py)",
        "五、前端可视化模块",
        "  5.1 仪表板布局",
        "  5.2 六大图表模块",
        "六、关键技术难点与解决方案",
        "七、项目总结",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(2)
        for run in p.runs:
            run.font.size = Pt(11)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════
    # 一、项目概述
    # ═══════════════════════════════════════════════════════════
    add_heading_styled(doc, "一、项目概述", 1)

    doc.add_paragraph(
        "本项目是一个全球气温变化可视化分析大屏系统，旨在通过交互式图表直观展示"
        "全球气候变暖的长期趋势、区域差异和季节波动特征。系统后端使用 Python FastAPI "
        "框架构建 RESTful API，前端使用 ECharts 5 图表库渲染六大可视化模块。"
    )

    doc.add_paragraph("项目核心功能包括：")
    features = [
        "全球年均温变化趋势（陆地 / 海陆综合 / 最高 / 最低四条曲线对比）",
        "全球温度距平分析（基于 1951-1980 基准期）",
        "月均温季节周期展示",
        "全球各国平均温度地图（按年份可交互选择）",
        "纬度带（热带 / 温带 / 寒带）温度变化趋势对比",
        "城市平均温度排名（交互式年份和数量选择）",
    ]
    for f in features:
        doc.add_paragraph(f, style='List Bullet')  # 项目符号列表

    # ═══════════════════════════════════════════════════════════
    # 二、数据采集与处理
    # ═══════════════════════════════════════════════════════════
    add_heading_styled(doc, "二、数据采集与处理", 1)

    # 2.1 数据来源
    add_heading_styled(doc, "2.1 数据来源介绍", 2)
    doc.add_paragraph("本项目的温度数据来源于两个权威数据集：")

    sources = [
        ("NASA GISTEMP v4（戈达德空间研究所地表温度分析第四版）",
         "提供全球月度 L-OTI（Land-Ocean Temperature Index，海陆温度指数）距平数据，"
         "以及基于 GHCN v4（全球历史气候网络第四版）的 27,958 个气象台站的均温记录。"
         "GISTEMP 是目前全球气候变化研究领域最权威的数据集之一，被 IPCC 报告广泛引用。"),
        ("Berkeley Earth 地表温度数据集",
         "提供 1750-2015 年的全球月均温、各国/各城市/各州的平均气温数据，"
         "包含温度不确定度和最高/最低温等详细字段。项目原有数据即来源于此数据集。"),
    ]
    for title_text, desc in sources:
        p = doc.add_paragraph()
        run = p.add_run(title_text + "：")
        run.font.bold = True  # 来源名称加粗
        p.add_run(desc)

    # 2.2 数据采集流程
    add_heading_styled(doc, "2.2 NASA GISTEMP 数据采集流程", 2)
    doc.add_paragraph(
        "为将项目数据从 2013/2015 年更新至最新（2024-2026 年），编写了 "
        "自动化的数据更新脚本 update_data.py，具体流程如下："
    )

    steps = [
        ("Step 1：下载全球月度距平数据",
         "通过 HTTP 请求获取 NASA GISTEMP 的 GLB.Ts+dSST.csv 文件，"
         "该文件包含 1880 年至今的全球海陆综合温度距平值（相对于 1951-1980 基准期）。"),
        ("Step 2：下载台站清单与温度数据",
         "下载 v4.temperature.inv.txt（27,958 个台站的位置、名称清单）和 "
         "v4.mean_GISS_homogenized.txt.gz（35MB 压缩文件，包含所有台站的月度均温记录，"
         "数据单位为 0.01°C）。"),
        ("Step 3：台站→城市匹配",
         "通过地理坐标邻近度算法，将 22,578 个有效台站匹配到已有城市数据中的 "
         "3,448 个城市（200km 半径范围内），同时使用 FIPS 国家代码映射表将台站聚合到 243 个国家。"),
        ("Step 4：距平→绝对温度转换",
         "利用已有数据中 1951-1980 年的 LandAndOcean 平均值（15.30°C）作为基准，"
         "将 GISTEMP 距平值转换为绝对温度：绝对温度 = 基准均值 + 距平值。"),
        ("Step 5：追加写入 CSV",
         "将新生成的数据行追加到对应的 CSV 文件中，保留原有数据不变。"
         "全球温度追加 2016-2026 年数据（124 行），城市/国家数据追加 2014-2024 年数据。"),
    ]
    for title_text, desc in steps:
        p = doc.add_paragraph()
        run = p.add_run(title_text + "：")
        run.font.bold = True
        p.add_run(desc)

    doc.add_paragraph(
        "由于 GHCN v4 地面台站的上报存在 1-2 年的延迟，2025-2026 年仅极少数台站有记录，"
        "因此城市/国家级数据只更新到 2024 年。全球均值数据使用统计插值模型，可更新至 2026 年 4 月。"
    )

    # 2.3 数据格式转换
    add_heading_styled(doc, "2.3 数据格式转换要点", 2)
    doc.add_paragraph(
        "GISTEMP 提供的台站数据格式为 Fortran I5 固定宽度格式（每字段 5 字符），"
        "温度单位是 0.01°C。脚本采用按空白分词（split）的稳健解析方式，过滤缺失值标记（-9999）"
        "和异常值（|raw_value| > 8000），确保数据质量。同时通过 FIPS 10-4 代码映射表将台站"
        "ID 前缀转换为标准国家名称。"
    )

    # 数据更新汇总表
    add_table_with_style(doc,
        ["数据集文件", "原有数据截止", "更新后截止", "新增行数", "数据来源"],
        [
            ["GlobalTemperatures.csv", "2015-12", "2026-04", "124 行", "GLB.Ts+dSST 全球距平"],
            ["ByCountry.csv", "2013-09", "2024-12", "20,706 行", "GHCN v4 台站按国聚合"],
            ["ByCity.csv", "2013-09", "2024-12", "112,258 行", "GHCN v4 台站→城市匹配"],
            ["ByMajorCity.csv", "2013-09", "2024-12", "9,394 行", "GHCN v4 台站→主要城市"],
        ],
        col_widths=[4.0, 2.5, 2.5, 2.5, 4.0],
    )

    # ═══════════════════════════════════════════════════════════
    # 三、系统架构设计
    # ═══════════════════════════════════════════════════════════
    add_heading_styled(doc, "三、系统架构设计", 1)

    add_heading_styled(doc, "3.1 整体架构", 2)
    doc.add_paragraph(
        "系统采用经典的前后端分离架构。后端基于 Python FastAPI 框架，负责从 CSV 数据集中"
        "读取、清洗、聚合温度数据，并通过 7 个 RESTful API 端点以 JSON 格式提供给前端。"
        "前端为纯静态 HTML + CSS + JavaScript 单页面应用，使用 ECharts 5 渲染图表。"
    )

    # ASCII 架构图 — 用等宽字体展示数据流向
    arch_text = (
        "┌──────────────────────────────────────────────────┐\n"
        "│             NASA GISTEMP v4 数据源                  │\n"
        "│  · GLB.Ts+dSST.csv (全球月度距平)                   │\n"
        "│  · GHCN v4 台站数据 (27,958 站)                     │\n"
        "└─────────────────┬────────────────────────────────┘\n"
        "                  │  scripts/update_data.py\n"
        "                  │  (下载 + 转换 + 追加写入)\n"
        "                  ▼\n"
        "┌──────────────────────────────────────────────────┐\n"
        "│           src/dataset/ (5 个 CSV 文件)             │\n"
        "│  · GlobalTemperatures.csv                          │\n"
        "│  · GlobalLandTemperaturesByCountry.csv              │\n"
        "│  · GlobalLandTemperaturesByCity.csv                 │\n"
        "│  · GlobalLandTemperaturesByMajorCity.csv            │\n"
        "│  · GlobalLandTemperaturesByState.csv                │\n"
        "└─────────────────┬────────────────────────────────┘\n"
        "                  │  src/data_engine.py (DataEngine)\n"
        "                  │  (pandas 读取 → 清洗 → 聚合 → 计算)\n"
        "                  ▼\n"
        "┌──────────────────────────────────────────────────┐\n"
        "│       src/app.py (FastAPI, 7 个 API 路由)         │\n"
        "│  · GET /api/global/annual                          │\n"
        "│  · GET /api/global/anomaly                         │\n"
        "│  · GET /api/global/monthly                         │\n"
        "│  · GET /api/global/seasonal-anomaly                │\n"
        "│  · GET /api/country/annual                         │\n"
        "│  · GET /api/city-temp                               │\n"
        "│  · GET /api/city/latband                            │\n"
        "└─────────────────┬────────────────────────────────┘\n"
        "                  │  JSON (HTTP Response)\n"
        "                  ▼\n"
        "┌──────────────────────────────────────────────────┐\n"
        "│      前端 (templates/index.html + dashboard.js)    │\n"
        "│  · 6 个 ECharts 5 图表实例 (深色主题)               │\n"
        "│  · 交互式年份筛选 / 地图滑块 / 城市搜索               │\n"
        "└──────────────────────────────────────────────────┘"
    )
    add_code_block(doc, arch_text)

    # 3.2 技术栈
    add_heading_styled(doc, "3.2 技术栈", 2)
    add_table_with_style(doc,
        ["层次", "技术", "版本", "用途"],
        [
            ["后端框架", "FastAPI", "0.136", "RESTful API 路由、请求校验"],
            ["ASGI 服务器", "Uvicorn", "0.46", "高性能异步 HTTP 服务"],
            ["数据处理", "Pandas", "3.0", "CSV 读取、分组聚合、距平计算"],
            ["HTTP 客户端", "Requests", "2.34", "下载 NASA GISTEMP 数据"],
            ["前端图表", "ECharts", "5.5", "六大交互式可视化模块"],
            ["前端样式", "CSS Grid", "—", "响应式双列仪表板布局"],
            ["文档生成", "python-docx", "1.2", "生成答辩 Word 文档"],
        ],
        col_widths=[2.5, 3.0, 1.5, 8.5],
    )

    # ═══════════════════════════════════════════════════════════
    # 四、后端模块说明
    # ═══════════════════════════════════════════════════════════
    add_heading_styled(doc, "四、后端模块说明", 1)

    # 4.1 app.py
    add_heading_styled(doc, "4.1 FastAPI 路由层 (app.py)", 2)
    doc.add_paragraph(
        "app.py 是系统的 Web 入口，负责初始化 FastAPI 应用、挂载静态文件目录、"
        "实例化 DataEngine 数据引擎，并定义全部 7 个 API 端点。"
    )
    doc.add_paragraph("关键设计要点：")
    key_points = [
        "使用 FastAPI Query 参数校验，自动生成 OpenAPI 文档",
        "统一异常处理：FileNotFoundError → 404，其他异常 → 500",
        "NaN 安全序列化：自定义 _nan_to_none() 函数将 pandas NaN 转换为 JSON null",
        "StaticFiles 挂载，使前端可直接引用 /static/ 下的 CSS、JS 和 GeoJSON 文件",
    ]
    for pt in key_points:
        doc.add_paragraph(pt, style='List Bullet')

    # 4.2 data_engine.py
    add_heading_styled(doc, "4.2 数据引擎层 (data_engine.py)", 2)
    doc.add_paragraph(
        "DataEngine 类封装了所有数据访问逻辑，将 CSV 文件的底层操作与上层 API 解耦。"
        "每个公开方法对应一个 API 端点，方法内部使用 pandas 完成数据加载、清洗和计算。"
    )

    add_table_with_style(doc,
        ["方法", "对应 API", "核心逻辑"],
        [
            ["get_global_annual()", "/api/global/annual", "按年分组求均值，保留海陆综合温的部分数据行"],
            ["get_global_anomaly()", "/api/global/anomaly", "计算 1951-1980 基准均值，逐年减基准得距平"],
            ["get_global_monthly()", "/api/global/monthly", "按月份分组求多年平均，得出 12 个月的季节周期"],
            ["get_seasonal_anomaly()", "/api/global/seasonal-anomaly", "计算各月距平后 pivot 为年×月矩阵"],
            ["get_country_annual()", "/api/country/annual", "按国家和年份筛选，降序排列"],
            ["get_city_temp_by_year()", "/api/city-temp", "按城市和年份筛选，取前 N 个最高温城市"],
            ["get_city_latband()", "/api/city/latband", "根据纬度解析热带/温带/寒带，分组聚合"],
        ],
        col_widths=[3.5, 3.0, 9.0],
    )

    doc.add_paragraph(
        "其中 _parse_latitude() 静态方法将 '57.05N'/'10.33E' 格式的经纬度字符串转为浮点数，"
        "_latband() 方法根据纬度绝对值分类：≤23.5° 为热带，≤66.5° 为温带，>66.5° 为寒带。"
    )

    # 4.3 update_data.py
    add_heading_styled(doc, "4.3 数据更新脚本 (update_data.py)", 2)
    doc.add_paragraph(
        "这是一个独立的命令行脚本（517 行），用于定期从 NASA GISTEMP 下载最新数据"
        "并自动追加到项目的 CSV 文件中。核心函数包括："
    )
    funcs = [
        ("update_global_temperatures()",
         "下载 GLB.Ts+dSST.csv，解析月度距平，计算 1951-1980 基线均值（15.30°C），"
         "将距平转为绝对温度后追加写入。额外处理：只保留 2016+ 年数据以避免重复。"),
        ("load_station_inventory() / load_station_temperatures()",
         "下载台站清单和 35MB 的台站温度数据。使用按空白分词（split）的解析方式替代固定宽度"
         "解析，提高数据读取的鲁棒性。温度值除以 100（单位 0.01°C），过滤缺失标记和异常值。"),
        ("update_city_like_csv()",
         "通过地理坐标邻近度算法（欧氏距离 × 111km/度，200km 阈值）将 22,578 个台站匹配到"
         "已有的 3,448 个城市，聚合各城市多台站均值后追加写入。"),
        ("update_country_csv()",
         "通过 FIPS 10-4 代码映射（173 个国家/地区的 2 字母代码 → 全名），将台站按国家聚合。"
         "限制只产出 2014-2024 年的数据（避免 2025+ 台站覆盖不足导致的异常值）。"),
    ]
    for title_text, desc in funcs:
        p = doc.add_paragraph()
        run = p.add_run(title_text + "：")
        run.font.bold = True
        run.font.size = Pt(9)
        p.add_run(desc)

    # ═══════════════════════════════════════════════════════════
    # 五、前端可视化模块
    # ═══════════════════════════════════════════════════════════
    add_heading_styled(doc, "五、前端可视化模块", 1)

    add_heading_styled(doc, "5.1 仪表板布局", 2)
    doc.add_paragraph(
        "前端采用深色主题（Dark Theme）的 CSS Grid 双列布局。页面顶部设有全局年份筛选栏，"
        "可调整起始/结束年份并一键应用到三个全局图表。各图表卡片右下角显示加载状态指示器。"
    )

    add_heading_styled(doc, "5.2 六大图表模块", 2)
    doc.add_paragraph(
        "系统共包含 6 个 ECharts 图表实例，均由 dashboard.js（约 480 行）管理："
    )

    add_table_with_style(doc,
        ["编号", "图表名称", "图表类型", "数据 API", "交互特性"],
        [
            ["1", "全球温度变化趋势", "多线图 (4线)", "/api/global/annual",
             "dataZoom 缩放滑块，图例切换"],
            ["2", "全球温度距平", "柱状图 (红/蓝)", "/api/global/anomaly",
             "正值红色、负值蓝色条件着色"],
            ["3", "月均温季节周期", "折线+面积", "/api/global/monthly",
             "年均温参考线标记"],
            ["4", "全球各国平均温度", "世界地图", "/api/country/annual",
             "年份滑块、缩放漫游、悬停提示"],
            ["5", "纬度带温度趋势", "三线对比", "/api/city/latband",
             "热带/温带/寒带三色对比"],
            ["6", "城市温度排名", "水平柱状图", "/api/city-temp",
             "年份输入、数量选择、颜色渐变"],
        ],
        col_widths=[1.0, 3.0, 2.5, 3.5, 5.5],
    )

    doc.add_paragraph(
        "前端使用国家名映射表（COUNTRY_NAME_MAP）将数据集中的国家名与 GeoJSON 地图文件中的"
        "国家名对齐（如 'United States' → 'United States of America'）。世界地图 GeoJSON "
        "从本地 /static/data/ 加载，并设置了 GitHub 原始文件作为备用数据源。"
    )

    # ═══════════════════════════════════════════════════════════
    # 六、关键技术难点与解决方案
    # ═══════════════════════════════════════════════════════════
    add_heading_styled(doc, "六、关键技术难点与解决方案", 1)

    challenges = [
        ("难点一：NASA GISTEMP 数据与原有数据格式不兼容",
         "GISTEMP 提供的是温度距平（anomaly），而原有 Berkeley Earth 数据是绝对温度。"
         "解决方案：利用已有数据中 1951-1980 年的 LandAndOcean 平均值（15.30°C）作为基准，"
         "通过公式「绝对温度 = 基准均值 + 距平值」完成转换。对于城市/国家数据，直接使用 "
         "GHCN v4 台站的绝对温度记录（0.01°C 单位），无需距平转换。"),

        ("难点二：台站数据的固定宽度格式解析",
         "GISTEMP homogenized 台站数据采用 Fortran I5 固定宽度格式，负数值会导致列对齐偏移。"
         "解决方案：放弃固定宽度解析，改用按空白分词（line.split()）的方式读取每行 13 个 token"
         "（1 个年份 + 12 个月份值），大幅提高解析鲁棒性。"),

        ("难点三：2025-2026 年台站数据严重不足",
         "GHCN v4 全球台站网络存在 1-2 年的数据上报延迟，2025+ 年仅极少数台站有记录。"
         "解决方案：在脚本中加入 max_year=2024 限制，城市和国家级数据只更新到 2024 年。"
         "全球均值数据（GLB.Ts+dSST）使用统计插值模型，可更新至当前月份（2026-04）。"
         "前端根据实际 API 返回动态调整年份滑块范围。"),

        ("难点四：Python 3.14 alpha 版本的依赖兼容性",
         "项目原要求 Python ≥ 3.14，但 NumPy/pandas 等核心库在 Python 3.14 alpha 上缺少"
         "预编译的 C 扩展 wheel，导致 DLL 加载失败（Segmentation Fault）。"
         "解决方案：将 requires-python 降级至 ≥ 3.13，使用 Python 3.13.3 稳定版，"
         "所有依赖库均可正常编译和运行。"),

        ("难点五：NaN 值的 JSON 序列化",
         "pandas 的 NaN 值无法被 Python 标准 json 库序列化，导致 API 返回 500 错误。"
         "解决方案：编写 _nan_to_none() 辅助函数，在返回响应前遍历列表，将 float('nan') "
         "替换为 None（JSON null）。前端 ECharts 原生支持 null 值的数据点（显示为断点）。"),

        ("难点六：台站→城市的地理匹配",
         "GISTEMP 的 27,958 个台站使用独立 ID 体系，不直接对应城市名称。"
         "解决方案：计算每个台站与每个已知城市的欧氏距离（经纬度差 × 111km/度），"
         "选择 200km 范围内最近的匹配。使用 O(n×m) 遍历后，1,302 个城市成功匹配到台站。"),
    ]

    for title_text, desc in challenges:
        add_heading_styled(doc, title_text, 2)
        doc.add_paragraph(desc)

    # ═══════════════════════════════════════════════════════════
    # 七、项目总结
    # ═══════════════════════════════════════════════════════════
    add_heading_styled(doc, "七、项目总结", 1)

    doc.add_paragraph("本项目实现了一个功能完整的全球气温可视化分析系统，具有以下特点：")

    summary_points = [
        "数据权威性：整合 NASA GISTEMP v4 和 Berkeley Earth 两大权威气候数据集，"
        "覆盖 1750 年至 2026 年的全球温度记录",
        "可视化全面性：6 个图表模块从趋势、距平、季节、地理、纬度带、城市排名六个维度切入",
        "交互友好性：深色主题大屏风格，支持年份筛选、地图缩放漫游、图表缩放等交互操作",
        "数据可更新性：编写的 update_data.py 脚本可定期从 NASA 下载最新数据并自动导入",
        "代码规范性：前后端分离、分层架构清晰、异常处理完善、代码风格统一",
    ]
    for pt in summary_points:
        doc.add_paragraph(pt, style='List Bullet')

    doc.add_paragraph(
        "通过本项目，可以清晰地观察到全球变暖的显著趋势：自 1880 年以来，全球海陆综合温度"
        "已上升约 1.3°C，且 2015 年后的升温速率明显加快。纬度带分析显示，极地地区的升温"
        "幅度远高于热带地区（北极放大效应）。这些观察与国际气候科学界的研究结论高度一致。"
    )

    # ═══════════════════════════════════════════════════════════
    # 附录：项目文件清单
    # ═══════════════════════════════════════════════════════════
    doc.add_page_break()
    add_heading_styled(doc, "附录：项目文件清单", 1)

    add_table_with_style(doc,
        ["文件", "行数", "功能说明"],
        [
            ["src/app.py", "~165", "FastAPI 应用入口，7 个 API 路由"],
            ["src/data_engine.py", "~211", "数据引擎，CSV 加载/清洗/聚合/计算"],
            ["src/templates/index.html", "~86", "仪表板 HTML 页面"],
            ["src/static/js/dashboard.js", "~475", "6 个 ECharts 图表 + 交互控制"],
            ["src/static/css/style.css", "~120", "深色主题响应式样式"],
            ["scripts/update_data.py", "~517", "NASA GISTEMP 数据下载与导入"],
            ["scripts/generate_report.py", "本项目", "生成此 Word 答辩文档"],
            ["tests/api_test.py", "~80", "6 个 pytest API 测试用例"],
        ],
        col_widths=[5.5, 1.5, 8.5],
    )

    # ═══════════════════════════════════════════════════════════
    # 保存文档
    # ═══════════════════════════════════════════════════════════
    output_path = Path(__file__).resolve().parent.parent / "答辩报告_全球气温可视化.docx"
    doc.save(str(output_path))
    print(f"✓ 文档已生成: {output_path}")


if __name__ == "__main__":
    main()
